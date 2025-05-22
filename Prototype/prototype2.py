'''
Prototype 3: Testing document reading as input
Checking the usability of taking data as input to identify agent evaluation
Results:
docx2txt works very well to get info from .docx, will work for gatherin data before preprocessing
Only works for one scentence, need an aggragate method.

'''
from docx import Document
import spacy
import docx2txt
import os
from sentence_transformers import SentenceTransformer
import numpy as np

def extract_text_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in full_text:
                    full_text.append(text)

    return '\n'.join(full_text)

# Example usage
file_path = r"C:\Users\Evanw\OneDrive\Documents\GitHub\Summer2025\Prototype\my-resume.docx"
  # Replace with the path to your .docx resume
# Check if the file exists
if os.path.exists(file_path):
    print(f"File found at: {file_path}")
    # Extract text from the file
    resume_text = docx2txt.process(file_path)
    if resume_text:
        print("Extracted Text:")
        print(resume_text)
    else:
        print("Failed to extract text.")
else:
    print(f"File not found at {file_path}")

# Print the extracted resume text

# Load spaCy for text processing
nlp = spacy.load('en_core_web_sm')

# Initialize sentence transformer model for embedding
transformer_model = SentenceTransformer('all-MiniLM-L6-v2')

# Define roles you want to extract
roles_of_interest = ["organization", "research aptitude", "problem solving"]

# Split the text into sentences (for a more granular approach)
doc = nlp(resume_text)
sentences = [sent.text for sent in doc.sents]

# Convert the sentences into embeddings
sentence_embeddings = transformer_model.encode(sentences)

# Define a function to match sentences with your roles of interest
def extract_roles_from_text(sentences, embeddings, roles):
    relevant_sentences = []
    
    # For each role, you can use keyword matching, or similarity scores to identify relevant sentences
    for role in roles:
        for sentence, embedding in zip(sentences, embeddings):
            # Calculate similarity between the sentence and the role (you can define this as needed)
            similarity_score = np.dot(embedding, transformer_model.encode([role])[0])
            
            # If the sentence matches the role (you can adjust the threshold)
            if similarity_score > 0.35:  # You can tweak the threshold value
                relevant_sentences.append((role, sentence, similarity_score))
    
    return relevant_sentences

# Extract roles from the resume text
extracted_info = extract_roles_from_text(sentences, sentence_embeddings, roles_of_interest)

# Print out the extracted roles and the corresponding sentences
for role, sentence, score in extracted_info:
    print(f"Role: {role}")
    print(f"Sentence: {sentence}")
    print(f"Similarity Score: {score}\n")