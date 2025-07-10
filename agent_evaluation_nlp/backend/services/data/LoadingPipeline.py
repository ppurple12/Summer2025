'''
Loads all different types of different files according to their extension
Should handle: .docx, .txt, .pdf, .xlsx
Identifies category of file (resume, survey, feeback/review, statistics or other)
Returns JSON format to simulate query results from Document Database
'''

import json
import os
import pdfplumber
import pandas as pd
from docx import Document
from pprint import pprint


# simple categorization for now - may become machine learning
# TIME PERMITTING
def categorize_document(text, file_path=None):
    # check file path for keywords
    if file_path:
        file_name = os.path.basename(file_path).lower()  # get file name
        
        if "resume" in file_name or "cv" in file_name:
            return "resume"
        elif "survey" in file_name:
            return "survey"
        elif "review" in file_name or "feedback" in file_name or "performance" in file_name:
            return "review"
        elif "statistics" in file_name or "stats" in file_name:
            return "stats"
        
    # put entire document together to check for keywords
    if isinstance(text, list):
        flattened = []
        for entry in text:
            if isinstance(entry, list):
                flattened.extend(entry)
            else:
                flattened.append(entry)
        text = " ".join(flattened)

    text = text.lower()
    
    #search for keywords
    if "resume" in text or "cv" in text:
        return "resume"
    elif "survey" in text:
        return "survey"
    elif "review" in text or "feedback" in text or "performance" in text:
        return "review"
    elif "statistics" in text or "stats" in text:
        return "stats"
    else:
        return "other"

#to properly seperate documents into key-value pairs
def parse_into_sections_from_blocks(blocks):
    sections, current_header, section_content = {}, None, []

    # creates blocks to have headers and following content in order for K-V pairs
    for block in blocks:
        if block.startswith("[HEADER]"):
            if current_header:
                sections[current_header] = "\n".join(section_content)
            current_header = block.replace("[HEADER] ", "")
            section_content = []
        else:
            section_content.append(block)
    
    if current_header:
        sections[current_header] = "\n".join(section_content)

    return sections

# reads .docx files
def load_docx(path):
    doc = Document(path)
    blocks, current_header = [], None

    # headers are identified and stored, texts are grouped accordingly
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        is_heading = para.style.name.lower().startswith("heading") or text.isupper()

        # add accordingly
        if is_heading:
            current_header = text  
            blocks.append(f"[HEADER] {current_header}")
        else:
            blocks.append(text)  

    # gather all table data
    table_content = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_content.append(row_text)
    
    if table_content:
        # if no header exists for the table, append it under "General"
        
        blocks.append(f"[HEADER] General")
        blocks.extend(table_content)

    
    parsed_sections = parse_into_sections_from_blocks(blocks) #make into proper K-V map
    category = categorize_document("\n".join(blocks), path) # find according category
    
    #return data in proper format
    return {
            "source": path,
            "category": category, 
            "text": parsed_sections
        }



# reads .pdf files
def load_pdf(path):
    try:
        with pdfplumber.open(path) as pdf:
            full_text = []

            # headers are identified and stored, texts are grouped accordingly
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    for line in page_text.strip().split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        if line.isupper():
                            full_text.append(f"[HEADER] {line}")
                        else:
                            full_text.append(line)

        text_str = " ".join(full_text)
        category = categorize_document(text_str, path) # categorize accordingly

        #return data in proper format
        return {
            "source": path,
            "category": category, 
            "text": text_str
        }
    # error check
    except Exception as e:
        print(f"Error reading PDF {path}: {e}")
        return {
            "source": path, 
            "category": [],  
            "text": ""
        }

# read .xlsx files
def load_excel(path):
    try:
        df = pd.read_excel(path)
        documents = []
        for _, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            category = categorize_document(row_text, path)
            documents.append({
                "source": path,
                "category": category,
                "text": row.to_dict()  # make K-V pairs with table
            })
        return documents
    
    #error check
    except Exception as e:
        print(f"Error reading Excel {path}: {e}")
        return []

# read .csv files
def load_csv(path):
    try:
        df = pd.read_csv(path)
        documents = []
        for _, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            category = categorize_document(row_text, path)
            documents.append({
                "source": path,
                "category": category,
                "text": row.to_dict() # make K-V pairs with table
            })
        return documents
    
    # error check
    except Exception as e:
        print(f"Error reading CSV {path}: {e}")
        return []

# for .txt files
def load_txt(path):
    try:
        # headers are identified and stored, texts are grouped accordingly
        with open(path, "r", encoding="utf-8") as f:
            content = f.read().strip().split("\n")
        full_text = []
        for line in content:
            line = line.strip()
            if line and line.isupper():
                full_text.append(f"[HEADER] {line}")
            else:
                full_text.append(line)

        text_str = " ".join(full_text)
        category = categorize_document(text_str, path)

        #return data in proper format
        return {
            "source": path,
            "category": category,
            "text": text_str
        }
    # error check
    except Exception as e:
        print(f"Error reading TXT {path}: {e}")
        return {
            "source": path,
            "category": [],
            "text": ""
        }

    
# Load all files based on file type 
def load_all_files(path):
    all_data = []
    for root, _, files in os.walk(path):
        for fname in files:
            file_path = os.path.join(root, fname)
            try:
                if fname.endswith(".docx"):
                    all_data.append(load_docx(file_path))
                elif fname.endswith(".pdf"):
                    all_data.append(load_pdf(file_path))
                elif fname.endswith(".xlsx"):
                    all_data.extend(load_excel(file_path))  # each row as one item
                elif fname.endswith(".csv"):
                    all_data.extend(load_csv(file_path))    # each row as one item
                elif fname.endswith(".txt"):
                    all_data.append(load_txt(file_path))
            except Exception as e:
                print(f"Failed to load {file_path}: {e}")
    
    return all_data

#pprint(load_all_files(r"C:\Users\Evanw\OneDrive\Documents\GitHub\Summer2025\Prototype\data\Inputs"))